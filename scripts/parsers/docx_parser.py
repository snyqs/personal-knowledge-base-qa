from docx import Document
from typing import List


class DOCXParser:
    def __init__(self):
        pass

    def parse(self, file_path: str) -> str:
        md_content = []
        try:
            doc = Document(file_path)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                style_name = para.style.name
                if style_name.startswith('Heading 1') or style_name == 'Title':
                    md_content.append(f"# {text}")
                elif style_name.startswith('Heading 2'):
                    md_content.append(f"## {text}")
                elif style_name.startswith('Heading 3'):
                    md_content.append(f"### {text}")
                elif style_name.startswith('Heading 4'):
                    md_content.append(f"#### {text}")
                else:
                    md_content.append(text)
            
            for table in doc.tables:
                md_table = self._table_to_markdown(table)
                if md_table:
                    md_content.append("\n")
                    md_content.append(md_table)
                    md_content.append("\n")
            
            return "\n".join(md_content)
        except Exception as e:
            raise RuntimeError(f"DOCX解析失败: {str(e)}")

    def _table_to_markdown(self, table) -> str:
        if not table.rows:
            return ""
        
        md_lines = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = cell_text.replace('\n', ' ')
                cells.append(cell_text)
            
            if row_idx == 0:
                md_lines.append(f"| {' | '.join(cells)} |")
                md_lines.append(f"| {' | '.join(['---'] * len(cells))} |")
            else:
                md_lines.append(f"| {' | '.join(cells)} |")
        
        return '\n'.join(md_lines)